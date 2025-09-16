from pathlib import Path

from loader import dp, bot, db
import logging

logging.basicConfig(level=logging.INFO)  # Yoki DEBUG
logger = logging.getLogger(__name__)

import os
import re
import requests
import logging
import asyncio
from datetime import datetime
from aiogram import Bot, Dispatcher, F, Router
from aiogram.types import Message, FSInputFile, InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery
from aiogram.filters import CommandStart, StateFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.client.default import DefaultBotProperties
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Configuration ---

AI_API_KEY = "sk-cb355fe140f84aad88999f6a3db45634"
AI_API_URL = "https://dashscope-intl.aliyuncs.com/compatible-mode/v1/chat/completions"


from aiogram import Router
from aiogram.types import Message
from aiogram.filters import Command

router = Router()  # Har doim shart!


# --- FSM States ---
class k_ReferatState(StatesGroup):
    k_doc_type = State()
    k_title = State()
    k_institute = State()
    k_author = State()
    k_pages = State()
    k_language = State()
    k_reja_mode = State()
    k_reja_manual = State()


# --- Keyboard functions ---








def k_language_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="О’zbek", callback_data="langkurs_uzbek"),
            InlineKeyboardButton(text="Русский", callback_data="langkurs_rus"),
            InlineKeyboardButton(text="English", callback_data="langkurs_english")
        ]
    ])


# --- AI communication function ---
async def k_ai_post(payload):
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: requests.post(AI_API_URL, headers={
        "Authorization": f"Bearer {AI_API_KEY}", "Content-Type": "application/json"
    }, json=payload))


# --- Outline generation function ---
async def k_generate_reja(title: str, lang: str = 'uzbek') -> list:
    lang_prompts = {
        "uzbek": f"'{title}' mavzusida faqat asosiy 3 ta banddan iborat reja yozing. Har bir bandni yangi qatorda raqamlangan ro'yxat shaklida yozing. Matn O'zbek tilida bo'lsin.",
        "rus": f"Напишите план по теме '{title}' из ровно 3 основных пунктов. Каждый пункт с новой строки в виде нумерованного списка. Текст на русском языке.",
        "english": f"Write exactly 3 main points for the topic '{title}'. Each point on a new line as a numbered list. Text should be in English."
    }
    prompt = lang_prompts.get(lang, lang_prompts["uzbek"])

    payload = {
        "model": "qwen-turbo",
        "messages": [
            {"role": "system",
             "content": "Siz reja tuzuvchi mutaxassis. Har doim faqat 3 ta asosiy band yozing. Qo'shimcha izoh yoki formatlash belgilarini ishlatmang."},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
        "max_tokens": 200,
    }
    try:
        response = await k_ai_post(payload)
        response.raise_for_status()
        content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")

        cleaned_reja = []
        for line in content.split("\n"):
            cleaned_line = re.sub(r'^[*-]?\s*\d+\.\s*', '', line).strip()
            if cleaned_line:
                cleaned_reja.append(f"{len(cleaned_reja) + 1}. {cleaned_line}")
            if len(cleaned_reja) == 3:  # faqat 3 ta band
                break

        if not cleaned_reja:
            default_reja = {
                "uzbek": ["1. Kirish", "2. Asosiy qism", "3. Xulosa"],
                "rus": ["1. Введение", "2. Основная часть", "3. Заключение"],
                "english": ["1. Introduction", "2. Main part", "3. Conclusion"]
            }
            return default_reja.get(lang, default_reja["uzbek"])
        return cleaned_reja
    except Exception as e:
        logger.error(f"Reja yaratishda xato: {e}")
        default_reja = {
            "uzbek": ["1. Kirish", "2. Asosiy qism", "3. Xulosa"],
            "rus": ["1. Введение", "2. Основная часть", "3. Заключение"],
            "english": ["1. Introduction", "2. Main part", "3. Conclusion"]
        }
        return default_reja.get(lang, default_reja["uzbek"])



# --- Text generation with page control ---
# --- Text generation with strict page control ---
import re
import math

def k_clean_text(text: str) -> str:
    # HTML teglarini olib tashlaymiz
    text = re.sub(r"<.*?>", "", text)
    # Maxsus belgilarni olib tashlaymiz
    text = re.sub(r'[*_~#+\-\=\[\]\{\}\(\)\|]', '', text)
    # Ortiqcha probellarni olib tashlaymiz
    text = re.sub(r'\s+', ' ', text).strip()
    return text


async def k_generate_coursework(title: str, reja: list, pages: int, lang: str = 'uzbek') -> str:
    words_per_page = 350
    total_words_target = words_per_page * pages

    num_sections = len(reja)
    words_per_section = total_words_target // num_sections if num_sections else total_words_target
    words_per_section = max(280, words_per_section)
    max_tokens_per_section = min(1500, words_per_section * 2)

    lang_prompts = {
        "uzbek": lambda item: f"'{item}' mavzusi haqida taxminan {words_per_section} so'zli oddiy matn yozing. Maxsus belgilar (** ## - *) ishlatilmasin.",
        "rus": lambda item: f"Напишите связанный текст на тему '{item}' на русском языке, примерно {words_per_section} слов. Не используйте специальные символы (** ## - *).",
        "english": lambda item: f"Write a coherent text about '{item}' in English, around {words_per_section} words. No special characters like **, ##, -, *."
    }

    prompt_generator = lang_prompts.get(lang, lang_prompts["uzbek"])
    cleaned_reja = [re.sub(r'^\d+\.\s*', '', item).strip() for item in reja]

    text = f"{title.upper()}\n\n"
    total_words_count = 0

    for item in cleaned_reja:
        prompt = prompt_generator(item)
        payload = {
            "model": "qwen-turbo",
            "messages": [
                {"role": "system", "content": f"Matn {words_per_section} so'z atrofida bo'lishi kerak. To'liq va tushunarli yozing."},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.7,
            "max_tokens": max_tokens_per_section
        }

        response = await k_ai_post(payload)
        content = k_clean_text(response.json().get("choices", [{}])[0].get("message", {}).get("content", ""))

        # Juda uzun bo‘lsa — qisqartiramiz
        if len(content.split()) > words_per_section:
            content = " ".join(content.split()[:words_per_section])

        total_words_count += len(content.split())
        text += f"\n\n{item.upper()}\n\n{content}"

    # Yakuniy so‘zlar sonini tekshirish
    all_words = text.split()
    if len(all_words) > total_words_target:
        text = " ".join(all_words[:total_words_target])

    return text



# --- Estimate page count ---
def k_estimate_page_count(text: str) -> int:
    word_count = len(re.findall(r'\w+', text))
    return max(1, round(word_count / 350))  # 350 words per page


# --- Document formatting functions ---
def k_add_double_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'double')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '18')
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)


def k_remove_borders(section):
    for b in section._sectPr.xpath('./w:pgBorders'):
        section._sectPr.remove(b)


def k_clean_text(text: str) -> str:
    cleaned = re.sub(r'[*_~#+\-\=\[\]\{\}\(\)\|]', '', text)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned


# --- Save to Word with accurate page control ---
def k_save_to_word(title, institute, author, pages, language, text, reja, doc_type):
    doc = Document()
    section1 = doc.sections[0]
    section1.top_margin = Inches(1)
    section1.bottom_margin = Inches(1)
    section1.left_margin = Inches(1)
    section1.right_margin = Inches(1)
    k_add_double_border(section1)

    translations = {
        "uzbek": {
            "ministry": "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM FAN VA INNOVATSIYA VAZIRLIGI",
            "topic": "Mavzu",
            "student": "O'quvchi",
            "academic_year": "o'quv yili",
            "plan": "Reja",
            "referat": "KURS ISHI",
            "mustaqil": "KURS ISHI"
        },
        "rus": {
            "ministry": "МИНИСТЕРСТВО ВЫСШЕГО ОБРАЗОВАНИЯ, НАУКИ И ИННОВАЦИЙ РЕСПУБЛИКИ УЗБЕКИСТАН",
            "topic": "Тема",
            "student": "Студент",
            "academic_year": "учебный год",
            "plan": "План",
            "referat": "КУРСОВАЯ РАБОТА",
            "mustaqil": "КУРСОВАЯ РАБОТА"
        },
        "english": {
            "ministry": "MINISTRY OF HIGHER EDUCATION, SCIENCE AND INNOVATION OF THE REPUBLIC OF UZBEKISTAN",
            "topic": "Topic",
            "student": "Student",
            "academic_year": "academic year",
            "plan": "Plan",
            "referat": "COURSE WORK",
            "mustaqil": "COURSE WORK"
        }
    }

    lang_texts = translations.get(language, translations["uzbek"])
    print(doc_type,'tupe2')
    translated_doc_type = lang_texts["referat"] if doc_type == "KURS ISHI" else lang_texts["referat"]

    def k_set_style(p, size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0):
        p.alignment = align
        run = p.runs[0]
        run.font.size = Pt(size)
        run.bold = bold
        p.paragraph_format.space_after = Pt(after)

    # Add document header
    p1 = doc.add_paragraph(lang_texts["ministry"])
    k_set_style(p1, 16, True, WD_ALIGN_PARAGRAPH.CENTER, 12)
    p2 = doc.add_paragraph(institute)
    k_set_style(p2, 16, True, WD_ALIGN_PARAGRAPH.CENTER, 60)
    p3 = doc.add_paragraph(translated_doc_type)
    k_set_style(p3, 54, True, WD_ALIGN_PARAGRAPH.CENTER, 24)
    p4 = doc.add_paragraph(f"{lang_texts['topic']}: {title}")
    k_set_style(p4, 16, True, WD_ALIGN_PARAGRAPH.CENTER, 60)
    p5 = doc.add_paragraph(f"{lang_texts['student']}: {author}")
    k_set_style(p5, 14, True, WD_ALIGN_PARAGRAPH.CENTER, 230)
    p6 = doc.add_paragraph(f"{datetime.now().year}-{lang_texts['academic_year']}")
    k_set_style(p6, 9, True, WD_ALIGN_PARAGRAPH.CENTER)

    # Add content sections
    section2 = doc.add_section()
    k_remove_borders(section2)

    # Add outline
    p7 = doc.add_paragraph(f"{lang_texts['plan']}:")
    k_set_style(p7, 14, True, WD_ALIGN_PARAGRAPH.LEFT)
    for item in reja:
        pi = doc.add_paragraph(item)
        k_set_style(pi, 14, False, WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

    # Add main text with proper formatting
    for para in text.split("\n\n"):
        # **...** formatdagi matnlarni o'tkazib yuboramiz
        if para.startswith("**") and para.endswith("**"):
            continue  # Bu matnni butunlay o'tkazib yuboramiz

        # <b>...</b> teglaridagi matnlarni sarlavha sifatida qo'shamiz
        elif para.startswith("<b>") and para.endswith("</b>"):
            heading_text = para[3:-4].strip()

            p = doc.add_paragraph(heading_text)
            p.runs[0].font.size = Pt(18)
            p.runs[0].bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Qolgan matnlarni oddiy formatda qo'shamiz
        else:
            text = para.strip()
            if text:  # Matn bo'sh bo'lmaganligini tekshiramiz
                # Agar hujjatda shu matn yo'q bo'lsa qo'shamiz
                if not any(text == p.text for p in doc.paragraphs):
                    p = doc.add_paragraph(text)
                    p.runs[0].font.size = Pt(16)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    filename = f"kursishi_{title.replace(' ', '_')[:50]}.docx"
    doc.save(filename)
    return filename




@router.callback_query(F.data == "kurs_ishi")
async def start_referat(callback: CallbackQuery, state: FSMContext):
    check=db.select_user(tg_user=callback.from_user.id)
    if int(check[3]) < 5000:
        await callback.message.answer(
            f"Xisobingizda yetarli mablag' yetarli emas 🙅!\nXisobingiz:💰 {check[3]}\n\nXisobni toldirish uchun /bye tugmasini bosing.")
    else:
        doc_type = "KURS ISHI"

        await state.update_data(doc_type=doc_type)
        await callback.message.answer("Mavzuni kiriting:")
        await state.set_state(k_ReferatState.k_title)


@router.message(StateFilter(k_ReferatState.k_title))
async def get_title(message: Message, state: FSMContext):
    await state.update_data(title=k_clean_text(message.text))
    await message.answer("Institut va kafedrani kiriting:")
    await state.set_state(k_ReferatState.k_institute)


@router.message(StateFilter(k_ReferatState.k_institute))
async def get_institute(message: Message, state: FSMContext):
    await state.update_data(institute=k_clean_text(message.text))
    await message.answer("Muallif (Ism Familiya) ni kiriting:")
    await state.set_state(k_ReferatState.k_author)


@router.message(StateFilter(k_ReferatState.k_author))
async def get_author(message: Message, state: FSMContext):
    await state.update_data(author=k_clean_text(message.text))
    await message.answer("Asosiy matn uchun sahifalar sonini kiriting eng kami 10:")
    await state.set_state(k_ReferatState.k_pages)


@router.message(StateFilter(k_ReferatState.k_pages))
async def get_pages(message: Message, state: FSMContext):

    if not message.text.isdigit():
        await message.answer("Iltimos, raqam kiriting eng kami 10:")
        return
    summa = db.select_summa()
    tek = int(message.text) * int(summa[1])
    pages = int(message.text)
    if not 10 <= pages :
        await message.answer("Iltimos, eng kami 10 gacha bo'lgan raqam kiriting:")
        return
    check = db.select_user(tg_user=message.from_user.id)
    if int(check[3]) < tek:
        await message.answer(
            f"Xisobingizda yetarli mablag' yetarli emas 🙅!\nXisobingiz:💰 {check[3]}\n\nXisobni toldirish uchun /bye tugmasini bosing.")
    else:
        user_balans=db.select_user(tg_user=message.from_user.id)
        summa=db.select_summa()
        puli=int(message.text)*int(summa[1])
        update=int(user_balans[3])-int(puli)
        db.update_user_balans(tg_user=message.from_user.id,balans=update)
        data = await state.get_data()
        await state.update_data(pages=pages)
        await message.answer("Qaysi tilda yozilsin?", reply_markup=k_language_keyboard())
        await state.set_state(k_ReferatState.k_language)


@router.callback_query(F.data.startswith("langkurs_"))
async def get_language(callback: CallbackQuery, state: FSMContext):
    lang_code = callback.data.split("_")[1]
    lang = {"uzbek": "uzbek", "rus": "rus", "english": "english"}.get(lang_code, "uzbek")
    await state.update_data(language=lang)
    data = await state.get_data()
    ins=data['institute']
    ism=data['author']
    page=data['pages']
    til=data['language']

    data = await state.get_data()

    reja = await k_generate_reja(data['title'], data['language'])
    await state.update_data(reja_items=reja, reja_type="ai")


    await k_generate_and_send_referat(callback.message, state)

def get_progress_bar(step: int, total_steps: int = 10) -> str:
    filled = "●" * step
    empty = "○" * (total_steps - step)
    return filled + empty




import asyncio

async def k_generate_and_send_referat(message: Message, state: FSMContext):
    data = await state.get_data()
    requested_pages = data['pages']
    lang = data['language']
    doc_type = data['doc_type']
    reja = data['reja_items']

    progress_msg = await message.answer("Coursework is being prepared...")

    try:
        total_steps = 15
        for step in range(1, total_steps + 1):
            bar = get_progress_bar(step, total_steps)
            eta = total_steps - step  # taxminiy vaqt sekundlarda
            await progress_msg.edit_text(
                f"🕒 Coursework is being prepared:\n{bar}  "
            )
            await asyncio.sleep(19)  # har 0.4s da yangilanadi, umumiy ~4s

        # Matn generatsiya qilish
        text = await k_generate_coursework(data['title'], reja, requested_pages, lang)

        actual_pages = k_estimate_page_count(text)
        if actual_pages > requested_pages * 1.2:
            await progress_msg.delete()
            await message.answer(
                f"⚠️ The page {actual_pages} was generated instead of the page {requested_pages} you requested. "
"To send a shortened version, send the /short command."
            )
            await state.update_data(long_text=text)
            return

        # Hujjatni saqlash
        file_path = k_save_to_word(
            data['title'], data['institute'], data['author'],
            requested_pages, lang, text, reja, doc_type
        )

        await progress_msg.delete()
        caption = f"📌 {data['title']}\n💡 @Slaydai_bot\n\n Press /start to perform another action."
        await message.answer_document(FSInputFile(file_path), caption=caption)
        os.remove(file_path)

    except Exception as e:
        logger.error(f"Xatolik: {e}")
        await message.answer(f"An error occurred: {str(e)}")
    finally:
        await state.clear()



@router.message(F.text == "/qisqa")
async def send_short_version(message: Message, state: FSMContext):
    data = await state.get_data()
    long_text = data.get('long_text', '')

    if not long_text:
        await message.answer("Qisqartiriladigan matn topilmadi.")
        return

    progress_msg = await message.answer("Qisqartirilgan versiya tayyorlanmoqda...")

    try:
        # Generate shorter version
        requested_pages = data['pages']
        text = await k_generate_coursework(data['title'], data['reja_items'], requested_pages, data['language'])

        # Save and send
        file_path = k_save_to_word(
            data['title'], data['institute'], data['author'],
            requested_pages, data['language'], text, data['reja_items'], data['doc_type']
        )

        actual_pages = k_estimate_page_count(text)
        await progress_msg.delete()
        caption = f"📌 {data['title']} (qisqartirilgan)\n sahifa\n💡 @Slaydai_bot\n\n Boshqa amal bajarish uchun /start tugmasini bosing."
        await message.answer_document(FSInputFile(file_path), caption=caption)
        os.remove(file_path)

    except Exception as e:
        logger.error(f"Xatolik: {e}")
        await message.answer(f"Qisqartirishda xatolik: {str(e)}")
    finally:
        await state.clear()


