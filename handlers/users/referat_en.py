from pathlib import Path
from aiogram.fsm.state import StatesGroup, State

from handlers.users.referat_shablon_en import RefesState_shablon_en_en
from loader import dp, bot, db
import logging
class RefesState_shablon_en(StatesGroup):
    doc_type = State()
    title = State()
    institute = State()
    author = State()
    pages = State()
    language = State()
    reja_mode = State()
    reja_manual = State()
    tanlov = State()
    rad_etish = State()
    tayyorlash = State()
    ozgartirish = State()
    mavzu = State()
    til = State()
    muallif = State()
    institutt = State()

logging.basicConfig(level=logging.INFO)  # Yoki DEBUG
logger = logging.getLogger(__name__)

import os
import re
import requests
import logging
import asyncio
from datetime import datetime
from aiogram import Bot, Dispatcher, F, Router
from aiogram.types import Message, FSInputFile, InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery
from aiogram.filters import CommandStart, StateFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.client.default import DefaultBotProperties
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Configuration ---

AI_API_KEY = "sk-cb355fe140f84aad88999f6a3db45634"
AI_API_URL = "https://dashscope-intl.aliyuncs.com/compatible-mode/v1/chat/completions"


from aiogram import Router
from aiogram.types import Message
from aiogram.filters import Command

router = Router()  # Har doim shart!


# --- FSM States ---
class ReferatState_en_en(StatesGroup):
    doc_type = State()
    title = State()
    institute = State()
    author = State()
    pages = State()
    language = State()
    reja_mode = State()
    reja_manual = State()
    tanlov=State()
    rad_etish=State()
    tayyorlash=State()
    ozgartirish=State()
    mavzu=State()
    til=State()
    muallif=State()
    institutt=State()

# --- Keyboard functions ---
def start_keyboard_en():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📘 Referat / Mustaqil ish", callback_data="start_referat")]
    ])


def doc_type_keyboard_en():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="REFERAT", callback_data="set_type_en_referat")],
        [InlineKeyboardButton(text="INDEPENDENT WORK", callback_data="-_en_mustaqil")]
    ])





def language_keyboard_en():
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="О’zbek", callback_data="langen_uzbek"),
            InlineKeyboardButton(text="Русский", callback_data="langen_rus"),
            InlineKeyboardButton(text="English", callback_data="langen_english")
        ]
    ])


# --- AI communication function ---
async def ai_post_en(payload):
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: requests.post(AI_API_URL, headers={
        "Authorization": f"Bearer {AI_API_KEY}", "Content-Type": "application/json"
    }, json=payload))


# --- Outline generation function ---
async def generate_reja_en(title: str, lang: str = 'uzbek') -> list:
    lang_prompts = {
        "uzbek": f"'{title}' mavzusida faqat asosiy bandlardan iborat reja yozing. Har bir bandni yangi qatorda va raqamlangan ro'yxat shaklida yozing. Matn O'zbek tilida bo'lsin.",
        "rus": f"Напишите план по теме '{title}'. Только основные пункты, каждый с новой строки в виде нумерованного списка. Текст должен быть на русском языке.",
        "english": f"Write an outline for the topic '{title}'. Only main points, each on a new line as a numbered list. Text should be in English."
    }
    prompt = lang_prompts.get(lang, lang_prompts["uzbek"])

    payload = {
        "model": "qwen-turbo",
        "messages": [
            {"role": "system",
             "content": "Siz reja tuzuvchi mutaxassis. Reja bandlarini faqat raqamlangan ro'yxat sifatida, hech qanday qo'shimcha formatlash yoki maxsus simvollarsiz taqdim eting."},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
        "max_tokens": 500,
    }
    try:
        response = await ai_post_en(payload)
        response.raise_for_status()
        content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")

        cleaned_reja = []
        for line in content.split("\n"):
            cleaned_line = re.sub(r'^[*-]?\s*\d+\.\s*', '', line).strip()
            cleaned_line = re.sub(r'[*#]', '', cleaned_line).strip()
            if cleaned_line:
                cleaned_reja.append(f"{len(cleaned_reja) + 1}. {cleaned_line}")

        if not cleaned_reja:
            default_reja = {
                "uzbek": ["1. Kirish", "2. Asosiy qism", "3. Xulosa"],
                "rus": ["1. Введение", "2. Основная часть", "3. Заключение"],
                "english": ["1. Introduction", "2. Main part", "3. Conclusion"]
            }
            return default_reja.get(lang, default_reja["uzbek"])
        return cleaned_reja
    except Exception as e:
        logger.error(f"Reja yaratishda xato: {e}")
        default_reja = {
            "uzbek": ["1. Kirish", "2. Asosiy qism", "3. Xulosa"],
            "rus": ["1. Введение", "2. Основная часть", "3. Заключение"],
            "english": ["1. Introduction", "2. Main part", "3. Conclusion"]
        }
        return default_reja.get(lang, default_reja["uzbek"])


# --- Text generation with page control ---
async def generate_coursework_en(title: str, reja: list, pages: int, lang: str = 'uzbek') -> str:
    # Aniqroq sahifa hisob-kitobi
    words_per_page = 280  # kamaytirildi, chunki Word formatida kam so'z kerak
    total_words_needed = words_per_page * pages
    num_sections = len(reja)

    # Har bir bo'lim uchun zarur so'zlar sonini hisoblash
    words_per_section = total_words_needed // num_sections

    # Qat'iy chegara - ortiqcha matn yaratmaslik uchun
    max_words_per_section = words_per_section

    # Token limitini aniq belgilash
    tokens_per_section = int(words_per_section * 1.3)  # kamaytirildi

    lang_prompts = {
        "uzbek": lambda item, word_count: (
            f"'{item}' mavzusi bo'yicha FAQAT {word_count} so'zdan iborat qisqa matn yozing. "
            f"Matn O'zbek tilida bo'lsin. Hech qanday belgilar ishlatmang. "
            f"MUHIM: {word_count} so'zdan ORTIQ yozmang!"
        ),
        "rus": lambda item, word_count: (
            f"Напишите КОРОТКИЙ текст на тему '{item}' СТРОГО в {word_count} слов на русском языке. "
            f"НЕ ПРЕВЫШАЙТЕ {word_count} слов! Без символов форматирования."
        ),
        "english": lambda item, word_count: (
            f"Write a SHORT text about '{item}' STRICTLY in {word_count} words in English. "
            f"DO NOT EXCEED {word_count} words! No formatting symbols."
        )
    }

    prompt_generator = lang_prompts.get(lang, lang_prompts["uzbek"])
    cleaned_reja = [re.sub(r'^\d+\.\s*', '', item).strip() for item in reja]

    text = ""
    total_generated_words = 0

    for i, item in enumerate(cleaned_reja):
        # Oxirgi bo'lim uchun qolgan so'zlarni hisobga olish
        if i == len(cleaned_reja) - 1:
            remaining_words = total_words_needed - total_generated_words
            section_word_count = max(150, remaining_words)  # min 150 so'z
        else:
            section_word_count = words_per_section

        text += f"\n\n<b>{item.upper()}</b>\n\n"
        prompt = prompt_generator(item, section_word_count)

        try:
            payload = {
                "model": "qwen-turbo",
                "messages": [
                    {"role": "system",
                     "content": f"Siz qisqa matn yozasiz. MAKSIMUM {section_word_count} so'z yozing. "
                                f"Ortiqcha yozmang! Qisqa va aniq bo'ling. So'zlar sonini sanab yozing."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.3,  # yana kamaytirildi
                "max_tokens": min(tokens_per_section, section_word_count * 2)  # qat'iy limit
            }

            response = await ai_post_en(payload)
            content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")

            # Matnni qat'iy qisqartirish
            content = clean_text_en(content)
            words = content.split()

            # Agar so'zlar soni ko'p bo'lsa, qat'iy qisqartirish
            if len(words) > section_word_count:
                content = ' '.join(words[:section_word_count])

            # So'zlar sonini hisoblash
            word_count = len(words[:section_word_count])
            total_generated_words += word_count

            text += content

            # Debug uchun log
            logger.info(f"Bo'lim '{item}': {word_count} so'z yaratildi (maqsad: {section_word_count})")

        except Exception as e:
            logger.error(f"Matn generatsiya qilishda xato: {e}")
            # Xatolik yuz berganda standart matn qo'shish
            default_text = f"Bu bo'limda {item} haqida batafsil ma'lumot beriladi. " * (section_word_count // 10)
            text += default_text
            total_generated_words += len(re.findall(r'\w+', default_text))

    # Jami yaratilgan so'zlar sonini log qilish
    final_word_count = len(re.findall(r'\w+', text))
    logger.info(f"Jami yaratilgan so'zlar: {final_word_count}, Maqsad: {total_words_needed}")

    return text


# --- Estimate page count ---
def estimate_page_count_en(text: str) -> int:
    # Aniqroq sahifa hisoblagichi - Word formatiga mos
    word_count = len(re.findall(r'\w+', text))
    # Word hujjatida 280-280 so'z = 1 sahifa (formatlash hisobga olingan holda)
    pages = word_count / 280
    return max(1, round(pages))


# --- Document formatting functions ---
def add_double_border_en(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'double')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '18')
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)


def remove_borders_en(section):
    for b in section._sectPr.xpath('./w:pgBorders'):
        section._sectPr.remove(b)


def clean_text_en(text: str) -> str:
    cleaned = re.sub(r'[*_~#+\-\=\[\]\{\}\(\)\|]', '', text)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned


# --- Save to Word with accurate page control ---
def save_to_word_en(title, institute, author, pages, language, text, reja, doc_type):
    doc = Document()
    section1 = doc.sections[0]

    # Sahifa o'lchamlarini standart qilish
    section1.top_margin = Inches(1)
    section1.bottom_margin = Inches(1)
    section1.left_margin = Inches(1.25)  # chapga biroz kengaytirildi
    section1.right_margin = Inches(1)

    add_double_border_en(section1)

    translations = {
        "uzbek": {
            "ministry": "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM FAN VA INNOVATSIYA VAZIRLIGI",
            "topic": "Mavzu",
            "student": "O'quvchi",
            "academic_year": "o'quv yili",
            "plan": "Reja",
            "referat": "REFERAT",
            "mustaqil": "MUSTAQIL ISH",
            "kurs_ishi": "KURS ISHI"
        },
        "rus": {
            "ministry": "МИНИСТЕРСТВО ВЫСШЕГО ОБРАЗОВАНИЯ, НАУКИ И ИННОВАЦИЙ РЕСПУБЛИКИ УЗБЕКИСТАН",
            "topic": "Тема",
            "student": "Студент",
            "academic_year": "учебный год",
            "plan": "План",
            "referat": "РЕФЕРАТ",
            "mustaqil": "САМОСТОЯТЕЛЬНАЯ РАБОТА",
            "kurs_ishi": "КУРСОВАЯ РАБОТА"
        },
        "english": {
            "ministry": "MINISTRY OF HIGHER EDUCATION, SCIENCE AND INNOVATION OF THE REPUBLIC OF UZBEKISTAN",
            "topic": "Topic",
            "student": "Student",
            "academic_year": "academic year",
            "plan": "Plan",
            "referat": "REFERAT",
            "mustaqil": "INDEPENDENT WORK",
            "kurs_ishi": "COURSE WORK"
        }
    }

    lang_texts = translations.get(language, translations["uzbek"])

    if doc_type == "REFERAT":
        translated_doc_type = lang_texts["referat"]
    elif doc_type == 'KURS ISHI':
        translated_doc_type = lang_texts["kurs_ishi"]
    else:
        translated_doc_type = lang_texts["mustaqil"]

    def set_style(p, size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0):
        p.alignment = align
        if p.runs:  # Runs mavjudligini tekshirish
            run = p.runs[0]
            run.font.size = Pt(size)
            run.font.name = "Times New Roman"
            run.bold = bold
        p.paragraph_format.space_after = Pt(after)

    # Muqova sahifasi
    p1 = doc.add_paragraph(lang_texts["ministry"])
    set_style(p1, 14, True, WD_ALIGN_PARAGRAPH.CENTER, 12)

    p2 = doc.add_paragraph(institute)
    set_style(p2, 16, True, WD_ALIGN_PARAGRAPH.CENTER, 60)

    p3 = doc.add_paragraph(translated_doc_type)
    set_style(p3, 20, True, WD_ALIGN_PARAGRAPH.CENTER, 18)  # o'lchamni kichraytirildi

    p4 = doc.add_paragraph(f"{lang_texts['topic']}: {title}")
    set_style(p4, 16, True, WD_ALIGN_PARAGRAPH.CENTER, 60)

    p5 = doc.add_paragraph(f"{lang_texts['student']}: {author}")
    set_style(p5, 14, True, WD_ALIGN_PARAGRAPH.CENTER, 200)

    p6 = doc.add_paragraph(f"{datetime.now().year}-{lang_texts['academic_year']}")
    set_style(p6, 12, True, WD_ALIGN_PARAGRAPH.CENTER)

    # Yangi sahifa - reja
    doc.add_page_break()
    section2 = doc.add_section()
    remove_borders_en(section2)

    # Rejani qo'shish
    p7 = doc.add_paragraph(f"{lang_texts['plan']}:")
    set_style(p7, 16, True, WD_ALIGN_PARAGRAPH.CENTER, 12)

    for item in reja:
        pi = doc.add_paragraph(item)
        set_style(pi, 14, False, WD_ALIGN_PARAGRAPH.LEFT, 6)

    doc.add_page_break()

    # Asosiy matnni qo'shish - aniqroq formatlash bilan
    paragraphs = text.split("\n\n")
    current_word_count = 0
    target_words = pages * 350  # maqsadli so'zlar soni

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Sarlavhalarni qayta ishlash
        if para.startswith("<b>") and para.endswith("</b>"):
            heading_text = para[3:-4].strip()
            if heading_text:
                p = doc.add_paragraph(heading_text)
                if p.runs:
                    p.runs[0].font.size = Pt(16)
                    p.runs[0].font.name = "Times New Roman"
                    p.runs[0].bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)

        # Oddiy matnni qayta ishlash
        else:
            # Formatni tozalash
            cleaned_para = clean_text_en(para)
            if cleaned_para and len(cleaned_para) > 10:  # juda qisqa matnlarni o'tkazib yuborish

                # So'zlar sonini tekshirish va kerak bo'lsa qisqartirish
                words_in_para = len(re.findall(r'\w+', cleaned_para))
                remaining_words = target_words - current_word_count

                if remaining_words > 0:
                    if words_in_para > remaining_words:
                        # Matnni qisqartirish
                        words = cleaned_para.split()
                        cleaned_para = ' '.join(words[:remaining_words])

                    p = doc.add_paragraph("\t" + cleaned_para)
                    if p.runs:
                        p.runs[0].font.size = Pt(14)
                        p.runs[0].font.name = "Times New Roman"
                        p.runs[0].bold = False
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.5  # qator orasini kengaytirish
                    p.paragraph_format.space_after = Pt(6)

                    current_word_count += len(re.findall(r'\w+', cleaned_para))

                # Maqsadli so'z soniga yetgan bo'lsa to'xtatish
                if current_word_count >= target_words:
                    break

    filename = f"referat_{title.replace(' ', '_')[:50]}.docx"
    doc.save(filename)

    # Yaratilgan hujjatdagi sahifalar sonini log qilish
    actual_word_count = current_word_count
    estimated_pages = max(1, round(actual_word_count / 350))
    logger.info(f"Yaratildi: {actual_word_count} so'z, Taxminiy {estimated_pages} sahifa")

    return filename


async def generate_targeted_content_en(title: str, reja: list, target_pages: int, lang: str) -> str:
    """
    Aniq sahifa soniga mos matn yaratish uchun maxsus funksiya
    """
    target_words = target_pages * 350
    sections_count = len(reja)
    words_per_section = target_words // sections_count

    # Har bir bo'lim uchun matn yaratish
    all_sections = []

    for i, item in enumerate(reja):
        cleaned_item = re.sub(r'^\d+\.\s*', '', item).strip()

        # Oxirgi bo'lim uchun qolgan so'zlarni hisoblash
        if i == sections_count - 1:
            remaining_words = target_words - sum(len(re.findall(r'\w+', section)) for section in all_sections)
            section_words = max(200, remaining_words)
        else:
            section_words = words_per_section

        # Matn yaratish uchun prompt
        prompts = {
            "uzbek": f"'{cleaned_item}' mavzusi bo'yicha aynan {section_words} so'zdan iborat "
                     f"akademik matn yozing. Matn O'zbek tilida, to'liq va mantiqiy bo'lishi kerak. "
                     f"Hech qanday maxsus belgilar ishlatmang.",

            "rus": f"Напишите академический текст на тему '{cleaned_item}' "
                   f"объемом ровно {section_words} слов на русском языке. "
                   f"Текст должен быть полным и логичным без специальных символов.",

            "english": f"Write an academic text about '{cleaned_item}' "
                       f"with exactly {section_words} words in English. "
                       f"The text should be complete and logical without special characters."
        }

        prompt = prompts.get(lang, prompts["uzbek"])

        try:
            # AI dan matn olish
            payload = {
                "model": "qwen-turbo",
                "messages": [
                    {"role": "system",
                     "content": f"Siz professional akademik yozuvchi siz. "
                                f"Aynan {section_words} so'z yarating, kam yoki ko'p emas."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": int(section_words * 1.8)  # so'z soniga mos token
            }

            response = await ai_post_en(payload)
            content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")

            # Matnni tozalash va so'zlar sonini tekshirish
            cleaned_content = clean_text_en(content)
            actual_words = len(re.findall(r'\w+', cleaned_content))

            # Agar so'zlar soni kam bo'lsa, qo'shimcha matn yaratish
            if actual_words < section_words * 0.8:  # 80% dan kam bo'lsa
                additional_prompt = prompts[lang].replace(
                    f"{section_words} so'z",
                    f"{section_words - actual_words} so'z qo'shimcha"
                )

                additional_payload = {
                    "model": "qwen-turbo",
                    "messages": [
                        {"role": "user", "content": additional_prompt}
                    ],
                    "temperature": 0.7,
                    "max_tokens": int((section_words - actual_words) * 1.8)
                }

                additional_response = await ai_post_en(additional_payload)
                additional_content = additional_response.json().get("choices", [{}])[0].get("message", {}).get(
                    "content", "")
                cleaned_content += " " + clean_text_en(additional_content)

            # So'zlar sonini maqsadli songa yaqinlashtirish
            words = cleaned_content.split()
            if len(words) > section_words:
                cleaned_content = ' '.join(words[:section_words])
            s=r'\w+'
            all_sections.append(cleaned_content)
            logger.info(f"'{cleaned_item}' bo'limi: {len(re.findall(s, cleaned_content))} so'z")

        except Exception as e:
            logger.error(f"Bo'lim yaratishda xato: {e}")
            # Standart matn
            default_section = f"Bu bo'limda {cleaned_item} haqida batafsil tahlil beriladi. " * (section_words // 15)
            all_sections.append(default_section)

    # Barcha bo'limlarni birlashtirish
    final_text = ""
    for i, (section_content, original_item) in enumerate(zip(all_sections, reja)):
        section_title = re.sub(r'^\d+\.\s*', '', original_item).strip()
        final_text += f"\n\n<b>{section_title.upper()}</b>\n\n{section_content}"

    # Yakuniy natijani log qilish
    final_word_count = len(re.findall(r'\w+', final_text))
    estimated_pages = round(final_word_count / 350)
    logger.info(f"Yakuniy natija: {final_word_count} so'z, {estimated_pages} sahifa")

    return final_text

@router.callback_query(F.data == "referat_en")
async def start_referat(callback: CallbackQuery, state: FSMContext):
    olish = db.select_shablon_kurs(tg_id=callback.from_user.id, til='en')
    await state.update_data({"from_id":str(callback.from_user.id)})
    if olish:

        pul = db.select_summa()
        summa = db.select_user(tg_user=callback.from_user.id)
        tekshir = int(pul[1]) * 5
        if int(summa[3]) >= tekshir:
            try:
                await callback.message.delete()
            except:
                pass

            await callback.message.answer("Enter the subject2:")
            await state.set_state(RefesState_shablon_en_en.title)

    else:
        pul = db.select_summa()
        summa = db.select_user(tg_user=callback.from_user.id)
        tekshir = int(pul[1]) * 5
        if int(summa[3]) >= tekshir:
            check=db.select_user(tg_user=callback.from_user.id)
            if int(check[3]) < 2500:
                await callback.message.answer(
                    f"You don't have enough funds in your account 🙅!\nYour account:💰 {check[3]}\n\nYou need /bye to top up your account.")

            else:

                await state.set_state(ReferatState_en_en.doc_type)
                photo_manzili = Path(r".\handlers\users\img_1.png")
                rasim = FSInputFile(photo_manzili)
                await bot.send_photo(chat_id=callback.from_user.id, caption="Choose:", photo=rasim,
                                     reply_markup=doc_type_keyboard_en())
        else:
            text = (f"💰Your balance: {pul[1]} soums\n\n"
                    f"Use the /bye command to top up your account.")
            await callback.message.answer(text, parse_mode="Markdown")

@router.callback_query(F.data.startswith("set_type_en_"))
async def set_doc_type(callback: CallbackQuery, state: FSMContext):
    doc_type = "REFERAT" if callback.data == "set_type_en_referat" else "MUSTAQIL ISH"
    await state.update_data(doc_type=doc_type)
    await callback.message.answer("Enter the subject:")
    await state.set_state(ReferatState_en_en.title)


@router.message(StateFilter(ReferatState_en_en.title))
async def get_title(message: Message, state: FSMContext):
    await state.update_data(title=clean_text_en(message.text))
    await message.answer("Enter the institute and department:")
    await state.set_state(ReferatState_en_en.institute)


@router.message(StateFilter(ReferatState_en_en.institute))
async def get_institute(message: Message, state: FSMContext):
    await state.update_data(institute=clean_text_en(message.text))
    await message.answer("Enter Author (First Name Last Name):")
    await state.set_state(ReferatState_en_en.author)


@router.message(StateFilter(ReferatState_en_en.author))
async def get_author(message: Message, state: FSMContext):
    await state.update_data(author=clean_text_en(message.text))
    await message.answer("Enter the number of pages for the main text (pages 5-20):")
    await state.set_state(ReferatState_en_en.pages)


@router.message(StateFilter(ReferatState_en_en.pages))
async def get_pages(message: Message, state: FSMContext):

    if not message.text.isdigit():
        await message.answer("Please enter a number only (10-50):")
        return
    summa = db.select_summa()
    tek = int(message.text) * int(summa[1])
    pages = int(message.text)
    if not 5 <= pages <= 20:
        await message.answer("Please enter pages 5 to 20.")
        return

    check=db.select_user(tg_user=message.from_user.id)
    if int(check[3]) < tek:
        await message.message.answer(
            f"You don't have enough funds in your account 🙅!\nYour account:💰 {check[3]}\n\nClick /bye to top up your account.")
    else:

        await state.update_data(pages=pages)
        await message.answer("In which language should it be written?", reply_markup=language_keyboard_en())
        await state.set_state(ReferatState_en_en.language)


@router.callback_query(F.data.startswith("langen_"))
async def get_language(callback: CallbackQuery, state: FSMContext):
    lang_code = callback.data.split("_")[1]
    lang = {"uzbek": "uzbek", "rus": "rus", "english": "english"}.get(lang_code, "uzbek")
    await state.update_data(language=lang)
    # shettan boshlanishi
    data = await state.get_data()
    mavzu = data.get("title")
    doc_type = data.get("doc_type")
    ism = data.get("author")
    ins = data.get("institute")
    til = data.get("language")
    requested_slide_count = data.get('pages')
    m_id=data.get('m_id')
    reja = await generate_reja_en(data['title'], data['language'])
    await state.update_data(reja_items=reja, reja_type="ai")
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="✅ Preparation", callback_data="tayyorlash"),
                InlineKeyboardButton(text="🚫 Rejection", callback_data="rad_etish"),
                InlineKeyboardButton(text="✏️ Change", callback_data="ozgartirish")
            ]
        ]
    )
    await bot.delete_message(chat_id=callback.from_user.id, message_id=m_id)
    a = await bot.send_message(chat_id=callback.from_user.id, text=
    "🌟Great, check the following information.\n\n"
f"{doc_type}\n"
f"Subject: {mavzu}\n"
f"Author: {ism}\n"
f"Institute and department: {ins}\n"

f"Number of pages: {requested_slide_count}\n"
f"Language: {til}\n\n"
f"• If the information is correct, click the '✅ Prepare' button!\n"
f"• Click the '✏️ Edit' button to change\n"
f"• Click the '🚫 Reject' button to cancel",
                               reply_markup=kb
                               )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState_en_en.tanlov)


@router.callback_query(F.data == "ozgartirish", ReferatState_en_en.tanlov)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    ol = await state.get_data()
    m_id = ol.get("m_id")
    uid = callback.from_user.id
    await bot.delete_message(chat_id=callback.from_user.id, message_id=m_id)
    data = await state.get_data()
    mavzu = data.get("title")
    doc_type = data.get("doc_type")
    ism = data.get("author")
    ins = data.get("institute")
    til = data.get("language")
    requested_slide_count = data.get('pages')
    reja = await generate_reja_en(data['title'], data['language'])
    await state.update_data(reja_items=reja, reja_type="ai")
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Subject", callback_data="mavzu"),
                InlineKeyboardButton(text="Author ", callback_data="muallif")

            ],
            [
                InlineKeyboardButton(text="Language", callback_data="til"),
                InlineKeyboardButton(text="Institute", callback_data="institut")
            ],
            [
                InlineKeyboardButton(text="-", callback_data="minus"),
                InlineKeyboardButton(text=f"Pages: {requested_slide_count}",
                                     callback_data=f"count:{requested_slide_count}"),
                InlineKeyboardButton(text="+", callback_data="plus")
            ],

            [
                InlineKeyboardButton(text="Save ✅", callback_data="saqlash"),
            ]
        ]
    )
    a = await bot.send_message(chat_id=callback.from_user.id, text=
    "🌟Great, check the following information.\n\n"
f"{doc_type}\n"
f"Subject: {mavzu}\n"
f"Author: {ism}\n"
f"Institute and Department: {ins}\n"

f"Number of Pages: {requested_slide_count}\n"
f"Language: {til}\n\n",
                               reply_markup=kb
                               )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState_en_en.ozgartirish)


@router.callback_query(F.data == "mavzu", ReferatState_en_en.ozgartirish)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    a = await callback.message.answer("Submit abstract topic:")
    await state.update_data({"mm_id": a.message_id})

    await state.set_state(ReferatState_en_en.mavzu)
    await callback.answer()
    print(2)


@router.callback_query(F.data == "til", ReferatState_en_en.ozgartirish)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🇺🇿 O'zbekcha", callback_data="uzbek")],
        [InlineKeyboardButton(text="🇷🇺 Русский", callback_data="rus")],
        [InlineKeyboardButton(text="🇬🇧 English", callback_data="english")]
    ])
    a = await callback.message.answer("Select presentation language:", reply_markup=keyboard)
    await state.update_data({"mm_id": a.message_id})
    await state.set_state(ReferatState_en_en.til)
    await callback.answer()
    print(1)


@router.callback_query(F.data == "muallif", ReferatState_en_en.ozgartirish)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    a = await callback.message.answer("Submit author name:")
    await state.update_data({"mm_id": a.message_id})
    await state.set_state(ReferatState_en_en.muallif)
    await callback.answer()
    print(3)


@router.callback_query(F.data == "institut", ReferatState_en_en.ozgartirish)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    a = await callback.message.answer("Enter the institute and department:")
    await state.update_data({"mm_id": a.message_id})
    await state.set_state(ReferatState_en_en.institutt)
    await callback.answer()
    print(4)


@router.callback_query(ReferatState_en_en.ozgartirish)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    data = callback.data

    if data == 'saqlash':

        ol = await state.get_data()
        m_id = ol.get("m_id")
        pages = ol.get('pages')
        await bot.delete_message(chat_id=callback.from_user.id, message_id=m_id)

        await generate_and_send_referat_enn(callback.message, state)
        user_balans = db.select_user(tg_user=callback.from_user.id)
        summa = db.select_summa()
        puli = int(pages) * int(summa[1])
        update = int(user_balans[3]) - int(puli)
        db.update_user_balans(tg_user=callback.from_user.id, balans=update)
        await state.clear()

    else:
        ol = await state.get_data()
        m_id = ol.get("m_id")
        mavzu = ol.get("title")
        doc_type = ol.get("doc_type")
        ism = ol.get("author")
        ins = ol.get("institute")
        til = ol.get("language")
        requested_slide_count = ol.get('pages')

        if data == "plus":
            if int(requested_slide_count) > 19:
                await callback.answer("You can select a maximum of 20 pages.❗️", show_alert=True)
            else:
                await state.update_data({"pages": int(
                    requested_slide_count) + 1})

        if data == "minus":
            if int(requested_slide_count) < 6:
                await callback.answer("You can select a maximum of 5 pages.❗️", show_alert=True)
            else:
                await state.update_data({"pages": int(
                    requested_slide_count) - 1})

        ol = await state.get_data()
        m_id = ol.get("m_id")
        mavzu = ol.get("title")
        doc_type = ol.get("doc_type")
        ism = ol.get("author")
        ins = ol.get("institute")
        til = ol.get("language")
        requested_slide_count = ol.get('pages')
        kb = InlineKeyboardMarkup(
            inline_keyboard=[
                [
                    InlineKeyboardButton(text="Subject", callback_data="mavzu"),
                    InlineKeyboardButton(text="Author ", callback_data="muallif")

                ],
                [
                    InlineKeyboardButton(text="Language", callback_data="til"),
                    InlineKeyboardButton(text="Institute", callback_data="institut")
                ],
                [
                    InlineKeyboardButton(text="-", callback_data="minus"),
                    InlineKeyboardButton(text=f"Pages: {requested_slide_count}",
                                         callback_data=f"count:{requested_slide_count}"),
                    InlineKeyboardButton(text="+", callback_data="plus")
                ],

                [
                    InlineKeyboardButton(text="Save ✅", callback_data="saqlash"),
                ]
            ]
        )
        a = await bot.edit_message_text(chat_id=callback.from_user.id, message_id=m_id, text=
        "🌟Great, check the following information.\n\n"
f"{doc_type}\n"
f"Subject: {mavzu}\n"
f"Author: {ism}\n"
f"Institute and Department: {ins}\n"

f"Number of Pages: {requested_slide_count}\n"
f"Language:{til}\n\n",
                                        reply_markup=kb
                                        )
        await state.update_data({"m_id": a.message_id})
        await state.set_state(ReferatState_en_en.ozgartirish)


@router.callback_query(F.data == "rad_etish", ReferatState_en_en.tanlov)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    await bot.send_message(chat_id=callback.from_user.id, text="Presentation has been canceled.\n"
"Click the 📕Presentation button below to create a new presentation!\n\n"
"Click /start to continue.")
    await state.clear()
    await callback.answer()


@router.callback_query(F.data == "reja_manual")
async def reja_manual(callback: CallbackQuery, state: FSMContext):
    await state.update_data(reja_type="manual")
    await callback.message.answer("Enter the plan items (write each on a new line):")
    await state.set_state(ReferatState_en_en.reja_manual)


@router.callback_query(F.data == "tayyorlash", ReferatState_en_en.tanlov)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    ol = await state.get_data()
    m_id = ol["m_id"]
    pages=ol.get('pages')
    await bot.delete_message(chat_id=callback.from_user.id, message_id=m_id)
    uid = callback.from_user.id
    user_balans = db.select_user(tg_user=callback.from_user.id)
    summa = db.select_summa()
    puli = int(pages) * int(summa[1])
    update = int(user_balans[3]) - int(puli)
    db.update_user_balans(tg_user=callback.from_user.id, balans=update)
    await state.update_data(pages=pages)
    await generate_and_send_referat_enn(callback.message, state)
    await state.clear()


@router.callback_query(StateFilter(ReferatState_en_en.til))
async def handle_mavzu_message(message: CallbackQuery, state: FSMContext):
    til = message.data
    ol = await state.get_data()
    mm_id = ol.get("mm_id")
    db.update_shablon_lang(tg_id=message.from_user.id, pre_tili='en', til=message.data)
    await bot.delete_message(chat_id=message.from_user.id, message_id=mm_id)
    await state.update_data(language=til)

    ol = await state.get_data()
    m_id = ol.get("m_id")
    mavzu = ol.get("title")
    doc_type = ol.get("doc_type")
    ism = ol.get("author")
    ins = ol.get("institute")
    til = ol.get("language")
    requested_slide_count = ol.get('pages')

    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Subject", callback_data="mavzu"),
                InlineKeyboardButton(text="Author ", callback_data="muallif")

            ],
            [
                InlineKeyboardButton(text="Language", callback_data="til"),
                InlineKeyboardButton(text="Institute", callback_data="institut")
            ],
            [
                InlineKeyboardButton(text="-", callback_data="minus"),
                InlineKeyboardButton(text=f"Pages: {requested_slide_count}",
                                     callback_data=f"count:{requested_slide_count}"),
                InlineKeyboardButton(text="+", callback_data="plus")
            ],

            [
                InlineKeyboardButton(text="Save ✅", callback_data="saqlash"),
            ]
        ]
    )
    a = await bot.edit_message_text(chat_id=message.from_user.id, message_id=m_id, text=
    "🌟Great, check the following information.\n\n"
f"{doc_type}\n"
f"Subject: {mavzu}\n"
f"Author: {ism}\n"
f"Institute and Department: {ins}\n"

f"Number of Pages: {requested_slide_count}\n"
f"Language: {til}\n\n",
                                    reply_markup=kb
                                    )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState_en_en.ozgartirish)


@router.message(StateFilter(ReferatState_en_en.muallif))
async def handle_mavzu_message(message: Message, state: FSMContext):
    ol = await state.get_data()
    mm_id = ol.get("mm_id")
    await bot.delete_message(chat_id=message.from_user.id, message_id=mm_id)
    await message.delete()
    await state.update_data(author=message.text)

    ol = await state.get_data()
    m_id = ol.get("m_id")
    mavzu = ol.get("title")
    doc_type = ol.get("doc_type")
    ism = ol.get("author")
    ins = ol.get("institute")
    til = ol.get("language")
    requested_slide_count = ol.get('pages')

    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Subject", callback_data="mavzu"),
                InlineKeyboardButton(text="Author ", callback_data="muallif")

            ],
            [
                InlineKeyboardButton(text="Language", callback_data="til"),
                InlineKeyboardButton(text="Institute", callback_data="institut")
            ],
            [
                InlineKeyboardButton(text="-", callback_data="minus"),
                InlineKeyboardButton(text=f"Pages: {requested_slide_count}",
                                     callback_data=f"count:{requested_slide_count}"),
                InlineKeyboardButton(text="+", callback_data="plus")
            ],

            [
                InlineKeyboardButton(text="Save ✅", callback_data="saqlash"),
            ]
        ]
    )
    a = await bot.edit_message_text(chat_id=message.from_user.id, message_id=m_id, text=
    "🌟Great, check the following information.\n\n"
f"{doc_type}\n"
f"Subject: {mavzu}\n"
f"Author: {ism}\n"
f"Institute and Department: {ins}\n"

f"Number of Pages: {requested_slide_count}\n"
f"Language: {til}\n\n",
                                    reply_markup=kb
                                    )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState_en_en.ozgartirish)


@router.message(StateFilter(ReferatState_en_en.mavzu))
async def handle_mavzu_message(message: Message, state: FSMContext):
    await state.update_data(title=message.text)
    await message.delete()

    ol = await state.get_data()
    m_id = ol.get("m_id")
    mm_id = ol.get("mm_id")
    mavzu = ol.get("title")
    doc_type = ol.get("doc_type")
    ism = ol.get("author")
    ins = ol.get("institute")
    til = ol.get("language")
    requested_slide_count = ol.get('pages')
    await bot.delete_message(chat_id=
                             message.from_user.id, message_id=mm_id)
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Subject", callback_data="mavzu"),
                InlineKeyboardButton(text="Author", callback_data="muallif")
            ],
            [
                InlineKeyboardButton(text="Language", callback_data="til"),
                InlineKeyboardButton(text="Institute", callback_data="institut")
            ],
            [
                InlineKeyboardButton(text="-", callback_data="minus"),
                InlineKeyboardButton(text=f"Pages: {requested_slide_count}",
                                     callback_data=f"count:{requested_slide_count}"),
                InlineKeyboardButton(text="+", callback_data="plus")
            ],
            [
                InlineKeyboardButton(text="Save ✅", callback_data="saqlash"),
            ]
        ]
    )

    message_text = (
        "🌟Great, check the following information.\n\n"
f"Abstract type: {doc_type}\n"
f"Subject: {mavzu}\n"
f"Author: {ism}\n"
f"Institute and department: {ins}\n"
f"Number of pages: {requested_slide_count}\n"
f"Language: {til}\n\n"
    )

    a = await bot.edit_message_text(
        chat_id=message.from_user.id,
        message_id=m_id,
        text=message_text,
        reply_markup=kb
    )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState_en_en.ozgartirish)


@router.message(StateFilter(ReferatState_en_en.institutt))
async def handle_mavzu_message(message: Message, state: FSMContext):
    ol = await state.get_data()
    mm_id = ol.get("mm_id")
    await bot.delete_message(chat_id=message.from_user.id, message_id=mm_id)
    await message.delete()
    await state.update_data(institute=message.text)

    ol = await state.get_data()
    m_id = ol.get("m_id")
    mavzu = ol.get("title")
    doc_type = ol.get("doc_type")
    ism = ol.get("author")
    ins = ol.get("institute")
    til = ol.get("language")
    requested_slide_count = ol.get('pages')

    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Subject", callback_data="mavzu"),
                InlineKeyboardButton(text="Author ", callback_data="muallif")

            ],
            [
                InlineKeyboardButton(text="Language", callback_data="til"),
                InlineKeyboardButton(text="Institute", callback_data="institut")
            ],
            [
                InlineKeyboardButton(text="-", callback_data="minus"),
                InlineKeyboardButton(text=f"Pages: {requested_slide_count}",
                                     callback_data=f"count:{requested_slide_count}"),
                InlineKeyboardButton(text="+", callback_data="plus")
            ],

            [
                InlineKeyboardButton(text="Save ✅", callback_data="saqlash"),
            ]
        ]
    )
    a = await bot.edit_message_text(chat_id=message.from_user.id, message_id=m_id, text=
    "🌟Great, check the following information.\n\n"
f"{doc_type}\n"
f"Subject: {mavzu}\n"
f"Author: {ism}\n"
f"Institute and Department: {ins}\n"

f"Number of Pages: {requested_slide_count}\n"
f"Language: {til}\n\n",
                                    reply_markup=kb
                                    )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState_en_en.ozgartirish)


# tugadi


@router.callback_query(F.data == "reja_manual")
async def reja_manual(callback: CallbackQuery, state: FSMContext):
    await state.update_data(reja_type="manual")
    await callback.message.answer("Enter the plan items (write each on a new line):")
    await state.set_state(ReferatState_en_en.reja_manual)


@router.message(StateFilter(ReferatState_en_en.reja_manual))
async def get_manual_reja(message: Message, state: FSMContext):
    reja = [f"{i + 1}. {clean_text_en(line)}" for i, line in enumerate(message.text.split("\n")) if clean_text_en(line)]
    if len(reja) < 3:
        await message.answer("Please include at least 3 plan items:")
        return
    await state.update_data(reja_items=reja)
    await generate_and_send_referat_enn(message, state)
    user_balans = db.select_user(tg_user=message.from_user.id)
    summa = db.select_summa()
    puli = int(message.text) * int(summa[1])
    update = int(user_balans[3]) - int(puli)
    db.update_user_balans(tg_user=message.from_user.id, balans=update)


def get_progress_bar_en(step: int, total_steps: int = 10) -> str:
    filled = "●" * step
    empty = "○" * (total_steps - step)
    return filled + empty




import asyncio


async def generate_and_send_referat_enn(message: Message, state: FSMContext):
    data = await state.get_data()
    requested_pages = int(data['pages'])-int(2)
    lang = data['language']
    doc_type = data['doc_type']
    reja = data['reja_items']
    doc_type = data.get("doc_type")
    ism = data.get("author")
    ins = data.get("institute")
    til = data.get("language")
    from_id = data.get("from_id")
    requested_slide_count = data.get('pages')
    db.add_shablon_kurs(tg_id=from_id, institut=ins, ism_fam=ism, sahifa_soni=requested_slide_count,
                        til='en', kurs_tili=til)
    progress_msg = await message.answer("The abstract is being prepared...")

    try:
        total_steps = 15
        for step in range(1, total_steps + 1):
            bar = get_progress_bar_en(step, total_steps)
            await progress_msg.edit_text(f"🕒 Preparing:\n{bar}")
            await asyncio.sleep(8)  # sekinroq progress

        # Maqsadli matn yaratish
        text = await generate_targeted_content_en(data['title'], reja, requested_pages, lang)

        # Sahifalar sonini tekshirish
        actual_pages = estimate_page_count_en(text)

        # Agar sahifalar soni juda farq qilsa, qayta yaratish
        page_difference = abs(actual_pages - requested_pages)
        if page_difference > 2:  # 2 sahifadan ortiq farq bo'lsa
            logger.warning(f"Page difference: {page_difference}. Recreating...")

            # Qayta yaratish
            if actual_pages > requested_pages:
                # Qisqartirish kerak
                adjustment_factor = requested_pages / actual_pages
                adjusted_pages = int(requested_pages * adjustment_factor)
                text = await generate_targeted_content_en(data['title'], reja, adjusted_pages, lang)
            else:
                # Ko'paytirish kerak
                additional_pages = requested_pages - actual_pages
                additional_text = await generate_targeted_content_en(data['title'], reja, additional_pages, lang)
                text += additional_text

        # Hujjatni saqlash
        file_path = save_to_word_en(
            data['title'], data['institute'], data['author'],
            requested_pages, lang, text, reja, doc_type
        )

        await progress_msg.delete()

        final_pages = estimate_page_count_en(text)
        caption = (f"📌 {data['title']}\n"
          
                   f"💡 @Slaydai_bot\n\n"
                   f"/start Press and hold to do another action.")

        await message.answer_document(FSInputFile(file_path), caption=caption)
        os.remove(file_path)

    except Exception as e:
        logger.error(f"Xatolik: {e}")
        await progress_msg.edit_text(f"Xatolik yuz berdi: {str(e)}")
    finally:
        await state.clear()


@router.message(F.text == "/qisqa")
async def send_short_version(message: Message, state: FSMContext):
    data = await state.get_data()
    long_text = data.get('long_text', '')

    if not long_text:
        await message.answer("No text to be shortened was found.")
        return

    progress_msg = await message.answer("A shortened version is being prepared...")

    try:
        # Generate shorter version
        requested_pages = data['pages']
        text = await generate_coursework_en(data['title'], data['reja_items'], requested_pages, data['language'])

        # Save and send
        file_path = save_to_word_en(
            data['title'], data['institute'], data['author'],
            requested_pages, data['language'], text, data['reja_items'], data['doc_type']
        )

        actual_pages = estimate_page_count_en(text)
        await progress_msg.delete()
        caption = f"📌 {data['title']} (shortened)\n page\n💡 @Slaydai_bot\n\n Press /start to perform another action."
        await message.answer_document(FSInputFile(file_path), caption=caption)
        os.remove(file_path)

    except Exception as e:
        logger.error(f"Xatolik: {e}")
        await message.answer(f"Error in abbreviation: {str(e)}")
    finally:
        await state.clear()


