
from pathlib import Path

from handlers.users.referat_shablon_uz import RefesState_shablon
from loader import dp, bot, db
import logging

logging.basicConfig(level=logging.INFO)  # Yoki DEBUG
logger = logging.getLogger(__name__)

import os
import re
import requests
import logging
import asyncio
from datetime import datetime
from aiogram import Bot, Dispatcher, F, Router
from aiogram.types import Message, FSInputFile, InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery
from aiogram.filters import CommandStart, StateFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.client.default import DefaultBotProperties
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Configuration ---

AI_API_KEY = "sk-cb355fe140f84aad88999f6a3db45634"
AI_API_URL = "https://dashscope-intl.aliyuncs.com/compatible-mode/v1/chat/completions"

from aiogram import Router
from aiogram.types import Message
from aiogram.filters import Command

router = Router()  # Har doim shart!


# --- FSM States ---
class ReferatState(StatesGroup):
    doc_type = State()
    title = State()
    institute = State()
    author = State()
    pages = State()
    language = State()
    reja_mode = State()
    reja_manual = State()
    tanlov = State()
    rad_etish = State()
    tayyorlash = State()
    ozgartirish = State()
    mavzu = State()
    til = State()
    muallif = State()
    institutt = State()


# --- Keyboard functions ---
def re_start_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📘 Referat / Mustaqil ish", callback_data="start_referat")]
    ])


def re_doc_type_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="REFERAT", callback_data="set_type_referat")],
        [InlineKeyboardButton(text="MUSTAQIL ISH", callback_data="set_type_mustaqil")]
    ])


def re_language_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [
            InlineKeyboardButton(text="О’zbek", callback_data="lang_uzbek"),
            InlineKeyboardButton(text="Русский", callback_data="lang_rus"),
            InlineKeyboardButton(text="English", callback_data="lang_english")
        ]
    ])


# --- AI communication function ---
async def re_ai_post(payload):
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: requests.post(AI_API_URL, headers={
        "Authorization": f"Bearer {AI_API_KEY}", "Content-Type": "application/json"
    }, json=payload))


# --- Outline generation function ---
async def re_generate_reja(title: str, pages: int, lang: str = 'uzbek') -> list:
    # Sahifalar soniga qarab reja bandlar sonini hisoblash
    # Minimal 3 band (kirish, asosiy, xulosa), maksimal sahifalarga mos ravishda oshirish
    # Har bir asosiy band uchun taxminan 2-3 sahifa, shuning uchun asosiy bandlar = (pages - 2) // 2 + 1
    num_main_sections = max(1, (pages - 2) // 2)  # Kirish va xulosa uchun 2 sahifa ajratib, qolganini asosiy bandlarga
    total_sections = num_main_sections + 2  # Kirish + asosiy(lar) + xulosa

    lang_prompts = {
        "uzbek": f"'{title}' mavzusida faqat {total_sections} ta asosiy banddan iborat reja yozing. Har bir bandni yangi qatorda va raqamlangan ro'yxat shaklida yozing. Matn O'zbek tilida bo'lsin.",
        "rus": f"Напишите план по теме '{title}'. Только {total_sections} основных пункта, каждый с новой строки в виде нумерованного списка. Текст должен быть на русском языке.",
        "english": f"Write an outline for the topic '{title}'. Only {total_sections} main points, each on a new line as a numbered list. Text should be in English."
    }

    prompt = lang_prompts.get(lang, lang_prompts["uzbek"])

    payload = {
        "model": "qwen-turbo",
        "messages": [
            {"role": "system",
             "content": "Siz reja tuzuvchi mutaxassis. Reja bandlarini faqat raqamlangan ro'yxat sifatida, hech qanday qo'shimcha formatlash yoki maxsus simvollarsiz taqdim eting."},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
        "max_tokens": 800,
    }
    try:
        response = await re_ai_post(payload)
        response.raise_for_status()
        content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")

        cleaned_reja = []
        for line in content.split("\n"):
            cleaned_line = re.sub(r'^[*-]?\s*\d+\.\s*', '', line).strip()
            cleaned_line = re.sub(r'[*#]', '', cleaned_line).strip()
            if cleaned_line:
                cleaned_reja.append(f"{len(cleaned_reja) + 1}. {cleaned_line}")

        # faqat kerakli sonini olish va kesish
        cleaned_reja = cleaned_reja[:total_sections]

        if not cleaned_reja:
            default_reja = {
                "uzbek": ["1. Kirish", "2. Asosiy qism", "3. Xulosa"],
                "rus": ["1. Введение", "2. Основная часть", "3. Заключение"],
                "english": ["1. Introduction", "2. Main part", "3. Conclusion"]
            }
            return default_reja.get(lang, default_reja["uzbek"])
        return cleaned_reja
    except Exception as e:
        logger.error(f"Reja yaratishda xato: {e}")
        default_reja = {
            "uzbek": ["1. Kirish", "2. Asosiy qism", "3. Xulosa"],
            "rus": ["1. Введение", "2. Основная часть", "3. Заключение"],
            "english": ["1. Introduction", "2. Main part", "3. Conclusion"]
        }
        return default_reja.get(lang, default_reja["uzbek"])


# --- Text generation with page control ---
# --- Text generation with page control ---
async def re_generate_coursework(title: str, reja: list, pages: int, lang: str = 'uzbek') -> str:
    # Bir sahifaga o'rtacha 300 so'z
    words_per_page = 300
    total_words_target = words_per_page * pages

    # === YANGI QO'SHIMCHA: foydalanuvchi so'ragan sahifa soniga mos reja bo'lishi ===
    num_sections = pages  # har bir sahifa uchun 1 ta bo‘lim qilamiz
    if len(reja) < num_sections:
        # agar reja qisqa bo‘lsa, avtomatik yangi bo‘lim qo‘shamiz
        for i in range(len(reja)+1, num_sections+1):
            reja.append(f"Bo'lim {i}")
    elif len(reja) > num_sections:
        # agar reja uzun bo‘lsa, kerakli qismigacha qisqartiramiz
        reja = reja[:num_sections]

    # Har bir bo'lim uchun so'zlar taqsimoti
    introduction_words = total_words_target * 0.15  # 15% kirish
    conclusion_words = total_words_target * 0.15    # 15% xulosa
    main_sections_words = total_words_target * 0.70 # 70% asosiy bo'limlar

    words_per_main_section = (
        main_sections_words // (num_sections - 2) if num_sections > 2 else main_sections_words
    )

    max_tokens_per_section = min(1500, words_per_main_section * 1.5)  # Token limiti

    lang_prompts = {
        "uzbek": lambda item, word_count: (
            f"{item} haqida to'liq va batafsil matn yozing. "
            f"Matn faqat O'zbek tilida bo'lsin, {int(word_count)} so'zdan oshmasin. "
            f"Matnda hech qanday belgilar (** ## - *) ishlatilmasin. "
            f"Faqat oddiy, silliq matn yozing."
        ),
        "rus": lambda item, word_count: (
            f"Напишите связанный и подробный текст на тему '{item}' на русском языке. "
            f"Не более {int(word_count)} слов. "
            f"Не используйте специальные символы (** ## - *). Только обычный текст без заголовков."
        ),
        "english": lambda item, word_count: (
            f"Write a coherent and detailed text about '{item}' in English. "
            f"No more than {int(word_count)} words. "
            f"Do not use any special characters like **, ##, -, or *. "
            f"Just plain paragraph text, no headings."
        )
    }

    prompt_generator = lang_prompts.get(lang, lang_prompts["uzbek"])
    cleaned_reja = [re.sub(r'^\d+\.\s*', '', item).strip() for item in reja]

    text = ""
    for i, item in enumerate(cleaned_reja):
        # So'zlar taqsimoti
        if i == 0:  # Kirish
            word_count = introduction_words
        elif i == len(cleaned_reja) - 1:  # Xulosa
            word_count = conclusion_words
        else:  # Asosiy bo'limlar
            word_count = words_per_main_section

        text += f"\n\n<b>{item.upper()}</b>\n\n"
        prompt = prompt_generator(item, word_count)

        try:
            payload = {
                "model": "qwen-turbo",
                "messages": [
                    {"role": "system",
                     "content": f"Matn {int(word_count)} so'zdan oshmasin. "
                                f"To'liq, batafsil va silliq yozing. "
                                f"Sarlavhalar yozmang."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.7,
                "max_tokens": int(max_tokens_per_section)
            }

            response = await re_ai_post(payload)
            content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")

            # Keraksiz belgilarni olib tashlash
            content = re.sub(r'^\s*\d+\.\s*', '', content)
            content = re.sub(r'^[A-Za-z\s]+\:\s*', '', content)
            content = re.sub(r'[*#-]', '', content)

            text += content
        except Exception as e:
            logger.error(f"Matn generatsiya qilishda xato: {e}")
            text += f"[{item} bo'limi uchun matn generatsiya qilinmadi]"

    return text


# --- Estimate page count ---
def re_estimate_page_count(text: str) -> int:
    word_count = len(re.findall(r'\w+', text))
    return max(1, round(word_count / 300))  # Endi 300 so'z per sahifa, generatsiya bilan moslashtirildi


# --- Document formatting functions ---
def re_add_double_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'double')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '18')
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)


def re_remove_borders(section):
    for b in section._sectPr.xpath('./w:pgBorders'):
        section._sectPr.remove(b)


def re_clean_text(text: str) -> str:
    cleaned = re.sub(r'[*_~#+\-\=\[\]\{\}\(\)\|]', '', text)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned


# --- Save to Word with accurate page control ---
def re_save_to_word(title, institute, author, pages, language, text, reja, doc_type):
    doc = Document()
    section1 = doc.sections[0]
    section1.top_margin = Inches(1)
    section1.bottom_margin = Inches(1)
    section1.left_margin = Inches(1)
    section1.right_margin = Inches(1)
    re_add_double_border(section1)

    translations = {
        "uzbek": {
            "ministry": "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM FAN VA INNOVATSIYA VAZIRLIGI",
            "topic": "Mavzu",
            "student": "O'quvchi",
            "academic_year": "o'quv yili",
            "plan": "Reja",
            "referat": "REFERAT",
            "mustaqil": "MUSTAQIL ISH",
            "kurs_ishi": "KURS ISHI"
        },
        "rus": {
            "ministry": "МИНИСТЕРСТВО ВЫСШЕГО ОБРАЗОВАНИЯ, НАУКИ И ИННОВАЦИЙ РЕСПУБЛИКИ УЗБЕКИСТАН",
            "topic": "Тема",
            "student": "Студент",
            "academic_year": "учебный год",
            "plan": "План",
            "referat": "РЕФЕРАТ",
            "mustaqil": "САМОСТОЯТЕЛЬНАЯ РАБОТА",
            "kurs_ishi": "КУРСОВАЯ РАБОТА"
        },
        "english": {
            "ministry": "MINISTRY OF HIGHER EDUCATION, SCIENCE AND INNOVATION OF THE REPUBLIC OF UZBEKISTAN",
            "topic": "Topic",
            "student": "Student",
            "academic_year": "academic year",
            "plan": "Plan",
            "referat": "REFERAT",
            "mustaqil": "INDEPENDENT WORK",
            "kurs_ishi": "COURSE WORK"
        }
    }

    lang_texts = translations.get(language, translations["uzbek"])
    print(doc_type, 'type1')
    if doc_type == "REFERAT":
        translated_doc_type = lang_texts["referat"]
    elif doc_type == 'KURS ISHI':
        translated_doc_type = lang_texts["kurs_ishi"]
    else:
        translated_doc_type = lang_texts["mustaqil"]

    def re_set_style(p, size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0):
        p.alignment = align
        run = p.runs[0]
        run.font.size = Pt(size)
        run.font.name = "Times new roman"
        run.bold = bold
        p.paragraph_format.space_after = Pt(after)

    # Add document header
    p1 = doc.add_paragraph(lang_texts["ministry"])
    re_set_style(p1, 16, True, WD_ALIGN_PARAGRAPH.CENTER, 12)
    p2 = doc.add_paragraph(institute)
    re_set_style(p2, 16, True, WD_ALIGN_PARAGRAPH.CENTER, 60)
    p3 = doc.add_paragraph(translated_doc_type)
    re_set_style(p3, 54, True, WD_ALIGN_PARAGRAPH.CENTER, 18)
    p4 = doc.add_paragraph(f"{lang_texts['topic']}: {title}")
    re_set_style(p4, 16, True, WD_ALIGN_PARAGRAPH.CENTER, 60)
    p5 = doc.add_paragraph(f"{lang_texts['student']}: {author}")
    re_set_style(p5, 14, True, WD_ALIGN_PARAGRAPH.CENTER, 230)
    p6 = doc.add_paragraph(f"{datetime.now().year}-{lang_texts['academic_year']}")
    re_set_style(p6, 9, True, WD_ALIGN_PARAGRAPH.CENTER)

    # Add content sections
    section2 = doc.add_section()
    re_remove_borders(section2)

    # Add outline
    p7 = doc.add_paragraph(f"{lang_texts['plan']}:")
    re_set_style(p7, 16, True, WD_ALIGN_PARAGRAPH.LEFT)
    for item in reja:
        pi = doc.add_paragraph(item)
        re_set_style(pi, 16, False, WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

    # Add main text with proper formatting
    for para in text.split("\n\n"):
        # **...** formatdagi matnlarni o'tkazib yuboramiz
        if para.startswith("**") and para.endswith("**"):
            continue  # Bu matnni butunlay o'tkazib yuboramiz

        # <b>...</b> teglaridagi matnlarni sarlavha sifatida qo'shamiz
        elif para.startswith("<b>") and para.endswith("</b>"):
            heading_text = para[3:-4].strip()

            p = doc.add_paragraph(heading_text)
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.name = "Times New Roman"
            p.runs[0].bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Sarlavhani o‘rtaga joylash


        # Qolgan matnlarni oddiy formatda qo'shamiz
        else:
            text = para.strip()
            if text:  # Matn bo'sh bo'lmaganligini tekshiramiz
                # Agar hujjatda shu matn yo'q bo'lsa qo'shamiz
                if not any(text == p.text for p in doc.paragraphs):
                    p = doc.add_paragraph("\t" + text)
                    p.runs[0].font.size = Pt(14)
                    p.runs[0].font.name = "Times new roman"
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    filename = f"referat_{title.replace(' ', '_')[:50]}.docx"
    doc.save(filename)
    return filename


@router.callback_query(F.data == "referat")
async def start_referat(callback: CallbackQuery, state: FSMContext):
    olish = db.select_shablon_kurs(tg_id=callback.from_user.id, til='uz')
    if olish:

        pul = db.select_summa()
        summa = db.select_user(tg_user=callback.from_user.id, til='uz')
        tekshir = int(pul[1]) * 5
        if int(summa[3]) >= tekshir:
            try:
                await callback.message.delete()
            except:
                pass

            await callback.message.answer("Mavzuni kiriting:")
            await state.set_state(RefesState_shablon.title)
        else:
            text = (f"💰Balansingiz: {pul[1]} so'm\n\n"
                    f"Xisobingizni to'ldirish uchun /bye buyrug'idan foydalaning.")
            await callback.message.answer(text, parse_mode="Markdown")

    else:
        pul = db.select_summa()
        summa = db.select_user(tg_user=callback.from_user.id)
        tekshir = int(pul[1]) * 5
        if int(summa[3]) >= tekshir:
            check = db.select_user(tg_user=callback.from_user.id)
            if int(check[3]) < 2500:
                await callback.message.answer(
                    f"Xisobingizda yetarli mablag' yetarli emas 🙅!\nXisobingiz:💰 {check[3]}\n\nXisobni toldirish uchun /bye tugmasini bosing.")

            else:

                await state.set_state(ReferatState.doc_type)
                photo_manzili = Path(r".\handlers\users\img_1.png")
                rasim = FSInputFile(photo_manzili)
                await bot.send_photo(chat_id=callback.from_user.id, caption="Tanlang:", photo=rasim,
                                     reply_markup=re_doc_type_keyboard())
        else:
            text = (f"💰Balansingiz: {pul[1]} so'm\n\n"
                    f"Xisobingizni to'ldirish uchun /bye buyrug'idan foydalaning.")
            await callback.message.answer(text, parse_mode="Markdown")


@router.callback_query(F.data.startswith("set_type_"))
async def set_doc_type(callback: CallbackQuery, state: FSMContext):
    doc_type = "REFERAT" if callback.data == "set_type_referat" else "MUSTAQIL ISH"
    await state.update_data(doc_type=doc_type)
    await callback.message.answer("Mavzuni kiriting:")
    await state.set_state(ReferatState.title)


@router.message(StateFilter(ReferatState.title))
async def get_title(message: Message, state: FSMContext):
    await state.update_data(title=re_clean_text(message.text))
    await message.answer("Institut va kafedrani kiriting:")
    await state.set_state(ReferatState.institute)


@router.message(StateFilter(ReferatState.institute))
async def get_institute(message: Message, state: FSMContext):
    await state.update_data(institute=re_clean_text(message.text))
    await message.answer("Muallif (Ism Familiya) ni kiriting:")
    await state.set_state(ReferatState.author)


@router.message(StateFilter(ReferatState.author))
async def get_author(message: Message, state: FSMContext):
    await state.update_data(author=re_clean_text(message.text))
    await message.answer("Asosiy matn uchun sahifalar sonini kiriting (5-20 bet):")
    await state.set_state(ReferatState.pages)


@router.message(StateFilter(ReferatState.pages))
async def get_pages(message: Message, state: FSMContext):
    if not message.text.isdigit():
        await message.answer("Iltimos, faqat raqam kiriting (10-50):")
        return
    summa = db.select_summa()
    tek = int(message.text) * int(summa[1])
    pages = int(message.text)
    if not 5 <= pages <= 20:
        await message.answer("Iltimos,  matn uchun sahifalar sonini kiriting (5-20 bet):")
        return

    check = db.select_user(tg_user=message.from_user.id)
    if int(check[3]) < tek:
        await message.message.answer(
            f"Xisobingizda yetarli mablag' yetarli emas 🙅!\nXisobingiz:💰 {check[3]}\n\nXisobni toldirish uchun /bye tugmasini bosing.")

    else:
        user_balans = db.select_user(tg_user=message.from_user.id)
        summa = db.select_summa()
        puli = int(message.text) * int(summa[1])
        update = int(user_balans[3]) - int(puli)
        db.update_user_balans(tg_user=message.from_user.id, balans=update)
        await state.update_data(pages=pages)
        await message.answer("Qaysi tilda yozilsin?", reply_markup=re_language_keyboard())
        await state.set_state(ReferatState.language)


@router.callback_query(F.data.startswith("lang_"))
async def get_language(callback: CallbackQuery, state: FSMContext):
    lang_code = callback.data.split("_")[1]
    lang = {"uzbek": "uzbek", "rus": "rus", "english": "english"}.get(lang_code, "uzbek")
    await state.update_data(language=lang)

    # shettan boshlanishi
    data = await state.get_data()
    mavzu = data.get("title")
    doc_type = data.get("doc_type")
    ism = data.get("author")
    ins = data.get("institute")
    til = data.get("language")
    requested_slide_count = data.get('pages')
    reja = await re_generate_reja(data['title'], data['pages'], data['language'])  # O'zgartirish: pages ni uzatish
    await state.update_data(reja_items=reja, reja_type="ai")
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="✅ Tayyorlash", callback_data="tayyorlash"),
                InlineKeyboardButton(text="🚫 Rad etish", callback_data="rad_etish"),
                InlineKeyboardButton(text="✏️ O'zgartirish", callback_data="ozgartirish")
            ]
        ]
    )

    a = await bot.send_message(chat_id=callback.from_user.id, text=
    "🌟Ajoyib, quyidagi ma'lumotlarni tekshiring.\n\n"
    f"{doc_type}\n"
    f"Mavzu: {mavzu}\n"
    f"Muallif: {ism}\n"
    f"Institut va kafedra: {ins}\n"

    f"Sahifalar soni: {requested_slide_count}\n"
    f"Tili: {til}\n\n"
    f"• Ma'lumotlar to'g'ri bo'lsa, '✅ Tayyorlash' tugmasini bosing!\n"
    f"• O'zgartirish uchun '✏️ O'zgartirish' tugmasini bosing\n"
    f"• Bekor qilish uchun '🚫 Rad etish' tugmasini bosing",
                               reply_markup=kb
                               )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState.tanlov)


@router.callback_query(F.data == "ozgartirish", ReferatState.tanlov)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    ol = await state.get_data()
    m_id = ol.get("m_id")
    uid = callback.from_user.id
    await bot.delete_message(chat_id=callback.from_user.id, message_id=m_id)
    data = await state.get_data()
    mavzu = data.get("title")
    doc_type = data.get("doc_type")
    ism = data.get("author")
    ins = data.get("institute")
    til = data.get("language")
    requested_slide_count = data.get('pages')
    reja = await re_generate_reja(data['title'], data['pages'], data['language'])  # O'zgartirish: pages ni uzatish
    await state.update_data(reja_items=reja, reja_type="ai")
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Mavzu", callback_data="mavzu"),
                InlineKeyboardButton(text="Muallif ", callback_data="muallif")

            ],
            [
                InlineKeyboardButton(text="Til", callback_data="til"),
                InlineKeyboardButton(text="Institut", callback_data="institut")
            ],
            [
                InlineKeyboardButton(text="-", callback_data="minus"),
                InlineKeyboardButton(text=f"Sahifalar: {requested_slide_count}",
                                     callback_data=f"count:{requested_slide_count}"),
                InlineKeyboardButton(text="+", callback_data="plus")
            ],

            [
                InlineKeyboardButton(text="Saqlash ✅", callback_data="saqlash"),
            ]
        ]
    )
    a = await bot.send_message(chat_id=callback.from_user.id, text=
    "🌟Ajoyib, quyidagi ma'lumotlarni tekshiring.\n\n"
    f"{doc_type}\n"
    f"Mavzu: {mavzu}\n"
    f"Muallif: {ism}\n"
    f"Institut va kafedra: {ins}\n"

    f"Sahifalar soni: {requested_slide_count}\n"
    f"Tili: {til}\n\n",
                               reply_markup=kb
                               )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState.ozgartirish)


@router.callback_query(F.data == "mavzu", ReferatState.ozgartirish)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    a = await callback.message.answer("Referat mavzusini yuboring:")
    await state.update_data({"mm_id": a.message_id})

    await state.set_state(ReferatState.mavzu)
    await callback.answer()
    print(2)


@router.callback_query(F.data == "til", ReferatState.ozgartirish)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🇺🇿 O'zbekcha", callback_data="uzbek")],
        [InlineKeyboardButton(text="🇷🇺 Русский", callback_data="rus")],
        [InlineKeyboardButton(text="🇬🇧 English", callback_data="english")]
    ])
    a = await callback.message.answer("Taqdimot tilini tanlang:", reply_markup=keyboard)
    await state.update_data({"mm_id": a.message_id})
    await state.set_state(ReferatState.til)
    await callback.answer()
    print(1)


@router.callback_query(F.data == "muallif", ReferatState.ozgartirish)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    a = await callback.message.answer("Muallif ismini yuboring:")
    await state.update_data({"mm_id": a.message_id})
    await state.set_state(ReferatState.muallif)
    await callback.answer()
    print(3)


@router.callback_query(F.data == "institut", ReferatState.ozgartirish)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    a = await callback.message.answer("Institut va kafedrani kiriting:")
    await state.update_data({"mm_id": a.message_id})
    await state.set_state(ReferatState.institutt)
    await callback.answer()
    print(4)


@router.callback_query(ReferatState.ozgartirish)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    data = callback.data

    if data == 'saqlash':

        ol = await state.get_data()
        m_id = ol.get("m_id")
        await bot.delete_message(chat_id=callback.from_user.id, message_id=m_id)

        await re_generate_and_send_referat(callback.message, state)

        await state.clear()

    else:
        ol = await state.get_data()
        m_id = ol.get("m_id")
        mavzu = ol.get("title")
        doc_type = ol.get("doc_type")
        ism = ol.get("author")
        ins = ol.get("institute")
        til = ol.get("language")
        requested_slide_count = ol.get('pages')

        if data == "plus":
            if int(requested_slide_count) > 19:
                await callback.answer("Eng kopi bilan 20 ta sahifa tanlashingiz mumkin.❗️", show_alert=True)
            else:
                await state.update_data({"pages": int(
                    requested_slide_count) + 1})

        if data == "minus":
            if int(requested_slide_count) < 6:
                await callback.answer("Eng kopi bilan 5 ta sahifa tanlashingiz mumkin.❗️", show_alert=True)
            else:
                await state.update_data({"pages": int(
                    requested_slide_count) - 1})

        ol = await state.get_data()
        m_id = ol.get("m_id")
        mavzu = ol.get("title")
        doc_type = ol.get("doc_type")
        ism = ol.get("author")
        ins = ol.get("institute")
        til = ol.get("language")
        requested_slide_count = ol.get('pages')
        kb = InlineKeyboardMarkup(
            inline_keyboard=[
                [
                    InlineKeyboardButton(text="Mavzu", callback_data="mavzu"),
                    InlineKeyboardButton(text="Muallif ", callback_data="muallif")

                ],
                [
                    InlineKeyboardButton(text="Til", callback_data="til"),
                    InlineKeyboardButton(text="Institut", callback_data="institut")
                ],
                [
                    InlineKeyboardButton(text="-", callback_data="minus"),
                    InlineKeyboardButton(text=f"Sahifalar: {requested_slide_count}",
                                         callback_data=f"count:{requested_slide_count}"),
                    InlineKeyboardButton(text="+", callback_data="plus")
                ],

                [
                    InlineKeyboardButton(text="Saqlash ✅", callback_data="saqlash"),
                ]
            ]
        )
        a = await bot.edit_message_text(chat_id=callback.from_user.id, message_id=m_id, text=
        "🌟Ajoyib, quyidagi ma'lumotlarni tekshiring.\n\n"
        f"{doc_type}\n"
        f"Mavzu: {mavzu}\n"
        f"Muallif: {ism}\n"
        f"Institut va kafedra: {ins}\n"

        f"Sahifalar soni: {requested_slide_count}\n"
        f"Tili: {til}\n\n",
                                        reply_markup=kb
                                        )
        await state.update_data({"m_id": a.message_id})
        await state.set_state(ReferatState.ozgartirish)


@router.callback_query(F.data == "rad_etish", ReferatState.tanlov)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    await bot.send_message(chat_id=callback.from_user.id, text="Taqdimot bekor qilindi.\n"
                                                               "Yangi taqdimot yaratish uchun quyidagi 📕Taqdimot tugmasini bosing!\n\n"
                                                               "Boshqa amallarni bajarish uchun /start tugmasini bosing.")

    await state.clear()
    await callback.answer()


@router.callback_query(F.data == "reja_manual")
async def reja_manual(callback: CallbackQuery, state: FSMContext):
    await state.update_data(reja_type="manual")
    await callback.message.answer("Reja bandlarini kiriting (har birini yangi qatorda yozing):")
    await state.set_state(ReferatState.reja_manual)


@router.callback_query(F.data == "tayyorlash", ReferatState.tanlov)
async def tayyorlash_handler(callback: CallbackQuery, state: FSMContext):
    ol = await state.get_data()
    m_id = ol["m_id"]
    await bot.delete_message(chat_id=callback.from_user.id, message_id=m_id)
    uid = callback.from_user.id
    await re_generate_and_send_referat(callback.message, state)
    await state.clear()


@router.callback_query(StateFilter(ReferatState.til))
async def handle_mavzu_message(message: CallbackQuery, state: FSMContext):
    til = message.data
    ol = await state.get_data()
    mm_id = ol.get("mm_id")
    db.update_shablon_lang(tg_id=message.from_user.id, pre_tili='uz', til=message.data)
    await bot.delete_message(chat_id=message.from_user.id, message_id=mm_id)
    await state.update_data(language=til)

    ol = await state.get_data()
    m_id = ol.get("m_id")
    mavzu = ol.get("title")
    doc_type = ol.get("doc_type")
    ism = ol.get("author")
    ins = ol.get("institute")
    til = ol.get("language")
    requested_slide_count = ol.get('pages')

    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Mavzu", callback_data="mavzu"),
                InlineKeyboardButton(text="Muallif ", callback_data="muallif")

            ],
            [
                InlineKeyboardButton(text="Til", callback_data="til"),
                InlineKeyboardButton(text="Institut", callback_data="institut")
            ],
            [
                InlineKeyboardButton(text="-", callback_data="minus"),
                InlineKeyboardButton(text=f"Sahifalar: {requested_slide_count}",
                                     callback_data=f"count:{requested_slide_count}"),
                InlineKeyboardButton(text="+", callback_data="plus")
            ],

            [
                InlineKeyboardButton(text="Saqlash ✅", callback_data="saqlash"),
            ]
        ]
    )
    a = await bot.edit_message_text(chat_id=message.from_user.id, message_id=m_id, text=
    "🌟Ajoyib, quyidagi ma'lumotlarni tekshiring.\n\n"
    f"{doc_type}\n"
    f"Mavzu: {mavzu}\n"
    f"Muallif: {ism}\n"
    f"Institut va kafedra: {ins}\n"

    f"Sahifalar soni: {requested_slide_count}\n"
    f"Tili: {til}\n\n",
                                    reply_markup=kb
                                    )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState.ozgartirish)


@router.message(StateFilter(ReferatState.muallif))
async def handle_mavzu_message(message: Message, state: FSMContext):
    ol = await state.get_data()
    mm_id = ol.get("mm_id")
    await bot.delete_message(chat_id=message.from_user.id, message_id=mm_id)
    await message.delete()
    await state.update_data(author=message.text)

    ol = await state.get_data()
    m_id = ol.get("m_id")
    mavzu = ol.get("title")
    doc_type = ol.get("doc_type")
    ism = ol.get("author")
    ins = ol.get("institute")
    til = ol.get("language")
    requested_slide_count = ol.get('pages')

    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Mavzu", callback_data="mavzu"),
                InlineKeyboardButton(text="Muallif ", callback_data="muallif")

            ],
            [
                InlineKeyboardButton(text="Til", callback_data="til"),
                InlineKeyboardButton(text="Institut", callback_data="institut")
            ],
            [
                InlineKeyboardButton(text="-", callback_data="minus"),
                InlineKeyboardButton(text=f"Sahifalar: {requested_slide_count}",
                                     callback_data=f"count:{requested_slide_count}"),
                InlineKeyboardButton(text="+", callback_data="plus")
            ],

            [
                InlineKeyboardButton(text="Saqlash ✅", callback_data="saqlash"),
            ]
        ]
    )
    a = await bot.edit_message_text(chat_id=message.from_user.id, message_id=m_id, text=
    "🌟Ajoyib, quyidagi ma'lumotlarni tekshiring.\n\n"
    f"{doc_type}\n"
    f"Mavzu: {mavzu}\n"
    f"Muallif: {ism}\n"
    f"Institut va kafedra: {ins}\n"

    f"Sahifalar soni: {requested_slide_count}\n"
    f"Tili: {til}\n\n",
                                    reply_markup=kb
                                    )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState.ozgartirish)


@router.message(StateFilter(ReferatState.mavzu))
async def handle_mavzu_message(message: Message, state: FSMContext):
    await state.update_data(title=message.text)
    await message.delete()

    ol = await state.get_data()
    m_id = ol.get("m_id")
    mm_id = ol.get("mm_id")
    mavzu = ol.get("title")
    doc_type = ol.get("doc_type")
    ism = ol.get("author")
    ins = ol.get("institute")
    til = ol.get("language")
    requested_slide_count = ol.get('pages')
    await bot.delete_message(chat_id=
                             message.from_user.id, message_id=mm_id)
    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Mavzu", callback_data="mavzu"),
                InlineKeyboardButton(text="Muallif", callback_data="muallif")
            ],
            [
                InlineKeyboardButton(text="Til", callback_data="til"),
                InlineKeyboardButton(text="Institut", callback_data="institut")
            ],
            [
                InlineKeyboardButton(text="-", callback_data="minus"),
                InlineKeyboardButton(text=f"Sahifalar: {requested_slide_count}",
                                     callback_data=f"count:{requested_slide_count}"),
                InlineKeyboardButton(text="+", callback_data="plus")
            ],
            [
                InlineKeyboardButton(text="Saqlash ✅", callback_data="saqlash"),
            ]
        ]
    )

    message_text = (
        "🌟Ajoyib, quyidagi ma'lumotlarni tekshiring.\n\n"
        f"Referat turi: {doc_type}\n"
        f"Mavzu: {mavzu}\n"
        f"Muallif: {ism}\n"
        f"Institut va kafedra: {ins}\n"
        f"Sahifalar soni: {requested_slide_count}\n"
        f"Tili: {til}\n\n"
    )

    a = await bot.edit_message_text(
        chat_id=message.from_user.id,
        message_id=m_id,
        text=message_text,
        reply_markup=kb
    )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState.ozgartirish)


@router.message(StateFilter(ReferatState.institutt))
async def handle_mavzu_message(message: Message, state: FSMContext):
    ol = await state.get_data()
    mm_id = ol.get("mm_id")
    await bot.delete_message(chat_id=message.from_user.id, message_id=mm_id)
    await message.delete()
    await state.update_data(institute=message.text)

    ol = await state.get_data()
    m_id = ol.get("m_id")
    mavzu = ol.get("title")
    doc_type = ol.get("doc_type")
    ism = ol.get("author")
    ins = ol.get("institute")
    til = ol.get("language")
    requested_slide_count = ol.get('pages')

    kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Mavzu", callback_data="mavzu"),
                InlineKeyboardButton(text="Muallif ", callback_data="muallif")

            ],
            [
                InlineKeyboardButton(text="Til", callback_data="til"),
                InlineKeyboardButton(text="Institut", callback_data="institut")
            ],
            [
                InlineKeyboardButton(text="-", callback_data="minus"),
                InlineKeyboardButton(text=f"Sahifalar: {requested_slide_count}",
                                     callback_data=f"count:{requested_slide_count}"),
                InlineKeyboardButton(text="+", callback_data="plus")
            ],

            [
                InlineKeyboardButton(text="Saqlash ✅", callback_data="saqlash"),
            ]
        ]
    )
    a = await bot.edit_message_text(chat_id=message.from_user.id, message_id=m_id, text=
    "🌟Ajoyib, quyidagi ma'lumotlarni tekshiring.\n\n"
    f"{doc_type}\n"
    f"Mavzu: {mavzu}\n"
    f"Muallif: {ism}\n"
    f"Institut va kafedra: {ins}\n"

    f"Sahifalar soni: {requested_slide_count}\n"
    f"Tili: {til}\n\n",
                                    reply_markup=kb
                                    )
    await state.update_data({"m_id": a.message_id})
    await state.set_state(ReferatState.ozgartirish)


# tugadi


def re_get_progress_bar(step: int, total_steps: int = 10) -> str:
    filled = "●" * step
    empty = "○" * (total_steps - step)
    return filled + empty


import asyncio


async def re_generate_and_send_referat(message: Message, state: FSMContext):
    data = await state.get_data()
    requested_pages = data['pages']
    lang = data['language']
    doc_type = data['doc_type']
    reja = data['reja_items']
    mavzu = data.get("title")
    doc_type = data.get("doc_type")
    ism = data.get("author")
    ins = data.get("institute")
    til = data.get("language")
    from_id = data.get("from_id")
    requested_slide_count = data.get('pages')
    db.add_shablon_kurs(tg_id=from_id, institut=ins, ism_fam=ism, sahifa_soni=requested_slide_count, til='uz',
                        kurs_tili=til)
    progress_msg = await message.answer("Tayyorlanmoqda...")

    try:
        total_steps = 15
        for step in range(1, total_steps + 1):
            bar = re_get_progress_bar(step, total_steps)
            eta = total_steps - step  # taxminiy vaqt sekundlarda
            await progress_msg.edit_text(
                f"🕒 Tayyorlanmoqda:\n{bar}"
            )
            await asyncio.sleep(10)  # har 0.4s da yangilanadi, umumiy ~4s

        # Matn generatsiya qilish
        text = await re_generate_coursework(data['title'], reja, requested_pages, lang)

        actual_pages = re_estimate_page_count(text)

        # Hujjatni saqlash
        file_path = re_save_to_word(
            data['title'], data['institute'], data['author'],
            requested_pages, lang, text, reja, doc_type
        )

        await progress_msg.delete()
        caption = f"📌 {data['title']}\n💡 @Slaydai_bot\n\n Boshqa amal bajarish uchun /start tugmasini bosing."
        await message.answer_document(FSInputFile(file_path), caption=caption)
        os.remove(file_path)


    except Exception as e:
        logger.error(f"Xatolik: {e}")
        await message.answer(f"Xatolik yuz berdi: {str(e)}")
    finally:
        await state.clear()


@router.message(F.text == "/qisqa")
async def send_short_version(message: Message, state: FSMContext):
    data = await state.get_data()
    long_text = data.get('long_text', '')

    if not long_text:
        await message.answer("Qisqartiriladigan matn topilmadi.")
        return

    progress_msg = await message.answer("Qisqartirilgan versiya tayyorlanmoqda...")

    try:
        # Generate shorter version
        requested_pages = data['pages']
        text = await re_generate_coursework(data['title'], data['reja_items'], requested_pages, data['language'])

        # Save and send
        file_path = re_save_to_word(
            data['title'], data['institute'], data['author'],
            requested_pages, data['language'], text, data['reja_items'], data['doc_type']
        )

        actual_pages = re_estimate_page_count(text)
        await progress_msg.delete()
        caption = f"📌 {data['title']} (qisqartirilgan)\n sahifa\n💡 @Slaydai_bot\n\n Boshqa amal bajarish uchun /start tugmasini bosing."
        await message.answer_document(FSInputFile(file_path), caption=caption)
        os.remove(file_path)

    except Exception as e:
        logger.error(f"Xatolik: {e}")
        await message.answer(f"Qisqartirishda xatolik: {str(e)}")
    finally:
        await state.clear()


